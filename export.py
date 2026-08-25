import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def clean_for_pdf(text):
    """
    Cleans text to make it compatible with FPDF standard fonts (Helvetica/Arial),
    removing or replacing characters that standard encodings cannot represent.
    """
    if not text:
        return ""
    # Map common smart/unicode punctuation to standard ASCII
    replacements = {
        '\u201c': '"', '\u201d': '"',  # Smart double quotes
        '\u2018': "'", '\u2019': "'",  # Smart single quotes
        '\u2014': '-', '\u2013': '-',  # Em and En dashes
        '\u2022': '*',                  # Bullet points
        '\u2026': '...',                # Ellipsis
        '\u00a0': ' ',                  # Non-breaking space
        '\u00e9': 'e', '\u00e1': 'a',   # Basic accent fallbacks
        '\u00ed': 'i', '\u00f3': 'o',
        '\u00fa': 'u', '\u00f1': 'n'
    }
    
    for uni_char, ascii_char in replacements.items():
        text = text.replace(uni_char, ascii_char)
        
    # Remove any other non-ASCII characters to prevent FPDF crash
    text = text.encode('ascii', errors='ignore').decode('ascii')
    return text

class SummarizationPDF(FPDF):
    """Custom PDF Layout Class inheriting from FPDF."""
    
    def header(self):
        # Header banner style
        self.set_fill_color(99, 102, 241) # Indigo Accent #6366f1
        self.rect(0, 0, 210, 30, 'F')
        
        self.set_text_color(255, 255, 255)
        self.set_font('Helvetica', 'B', 16)
        self.cell(0, 10, 'NLP TEXT SUMMARIZATION REPORT', ln=True, align='C')
        self.set_font('Helvetica', 'I', 9)
        self.cell(0, 5, 'AI-Engineered Extractive Document Summary', ln=True, align='C')
        self.ln(10)

    def footer(self):
        # Go to 1.5 cm from bottom
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(148, 163, 184) # slate-400
        # Page number
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')


def export_to_pdf(summary_text, metadata, output_filepath):
    """
    Generates a professional PDF document.
    """
    try:
        # Clean the text for PDF compatibility
        cleaned_summary = clean_for_pdf(summary_text)
        
        pdf = SummarizationPDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        
        # Meta Card Section
        pdf.set_y(35)
        pdf.set_fill_color(248, 250, 252) # slate-50
        pdf.set_draw_color(226, 232, 240) # slate-200
        pdf.rect(10, 35, 190, 35, 'FD')
        
        pdf.set_text_color(71, 85, 105) # slate-600
        pdf.set_font('Helvetica', 'B', 10)
        
        # Grid layout for metadata
        pdf.set_xy(12, 38)
        pdf.cell(90, 6, f"ALGORITHM: {metadata.get('algorithm', 'N/A').upper()}", ln=False)
        pdf.cell(90, 6, f"COMPRESSION RATIO: {metadata.get('compression_ratio', 'N/A')}", ln=True)
        
        pdf.set_x(12)
        pdf.cell(90, 6, f"ORIGINAL WORDS: {metadata.get('orig_words', 0)}", ln=False)
        pdf.cell(90, 6, f"SUMMARY WORDS: {metadata.get('sum_words', 0)}", ln=True)
        
        pdf.set_x(12)
        pdf.cell(90, 6, f"ORIGINAL SENTENCES: {metadata.get('orig_sents', 0)}", ln=False)
        pdf.cell(90, 6, f"SUMMARY SENTENCES: {metadata.get('sum_sents', 0)}", ln=True)

        pdf.set_x(12)
        pdf.cell(180, 6, f"GENERATED AT: {metadata.get('timestamp', 'N/A')}", ln=True)
        
        pdf.ln(12)
        
        # Summary Header
        pdf.set_text_color(30, 41, 59) # slate-800
        pdf.set_font('Helvetica', 'B', 13)
        pdf.cell(0, 8, 'Summary Details', ln=True)
        
        # Decorative divider line
        pdf.set_draw_color(99, 102, 241)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
        
        # Summary Text Body
        pdf.set_text_color(51, 65, 85) # slate-700
        pdf.set_font('Helvetica', '', 10.5)
        pdf.multi_cell(0, 6.5, cleaned_summary)
        
        # Keywords Section
        keywords = metadata.get('keywords', [])
        if keywords:
            pdf.ln(8)
            pdf.set_font('Helvetica', 'B', 11)
            pdf.cell(0, 8, 'Key Extract Terms', ln=True)
            pdf.set_draw_color(226, 232, 240)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)
            
            pdf.set_font('Helvetica', '', 9.5)
            kw_str = ", ".join(keywords)
            pdf.multi_cell(0, 5, kw_str)
            
        # Ensure directories exist
        os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
        pdf.output(output_filepath)
        return output_filepath
    except Exception as e:
        raise RuntimeError(f"Failed to generate PDF document: {e}")


def export_to_docx(summary_text, metadata, output_filepath):
    """
    Generates a beautifully styled Word Document using python-docx.
    """
    try:
        doc = docx.Document()
        
        # Document Title
        title = doc.add_paragraph()
        title_run = title.add_run("NLP Text Summarization Report")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = docx.shared.RGBColor(99, 102, 241) # Indigo #6366f1
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run("Generated by NLP Extractive Engine")
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = docx.shared.RGBColor(148, 163, 184) # Slate-400
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # Metadata Table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Shading Accent 1'
        
        # Populate table
        data = [
            ("Algorithm", metadata.get('algorithm', 'N/A').upper()),
            ("Compression Ratio", metadata.get('compression_ratio', 'N/A')),
            ("Word Reduction", f"Original: {metadata.get('orig_words', 0)} -> Summary: {metadata.get('sum_words', 0)}"),
            ("Sentence Count", f"Original: {metadata.get('orig_sents', 0)} -> Summary: {metadata.get('sum_sents', 0)}")
        ]
        
        for i, (label, val) in enumerate(data):
            row = table.rows[i]
            row.cells[0].paragraphs[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].text = str(val)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(18)
        
        # Summary Header
        h_summary = doc.add_paragraph()
        h_summary_run = h_summary.add_run("Extracted Summary")
        h_summary_run.font.name = 'Arial'
        h_summary_run.font.size = Pt(14)
        h_summary_run.font.bold = True
        h_summary_run.font.color.rgb = docx.shared.RGBColor(30, 41, 59) # Slate-800
        
        # Divider line approximation
        p_div = doc.add_paragraph()
        p_div_run = p_div.add_run("―" * 55)
        p_div_run.font.color.rgb = docx.shared.RGBColor(99, 102, 241)
        p_div.paragraph_format.space_after = Pt(6)
        
        # Summary Body
        p_body = doc.add_paragraph()
        p_body_run = p_body.add_run(summary_text)
        p_body_run.font.name = 'Calibri'
        p_body_run.font.size = Pt(11)
        p_body.paragraph_format.line_spacing = 1.15
        
        # Keywords
        keywords = metadata.get('keywords', [])
        if keywords:
            doc.add_paragraph().paragraph_format.space_before = Pt(18)
            h_kw = doc.add_paragraph()
            h_kw_run = h_kw.add_run("Key Terms")
            h_kw_run.font.name = 'Arial'
            h_kw_run.font.size = Pt(12)
            h_kw_run.font.bold = True
            
            p_kw = doc.add_paragraph()
            p_kw_run = p_kw.add_run(", ".join(keywords))
            p_kw_run.font.italic = True
            p_kw_run.font.size = Pt(10)
            
        # Ensure directories exist
        os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
        doc.save(output_filepath)
        return output_filepath
    except Exception as e:
        raise RuntimeError(f"Failed to generate DOCX document: {e}")


def export_to_txt(summary_text, metadata, output_filepath):
    """
    Generates a simple, clean raw text file summary report.
    """
    try:
        os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.write("=" * 60 + "\n")
            f.write("        NLP TEXT SUMMARIZATION REPORT\n")
            f.write("=" * 60 + "\n\n")
            f.write(f"Algorithm: {metadata.get('algorithm', 'N/A').upper()}\n")
            f.write(f"Compression Ratio: {metadata.get('compression_ratio', 'N/A')}\n")
            f.write(f"Original Words: {metadata.get('orig_words', 0)} | Summary Words: {metadata.get('sum_words', 0)}\n")
            f.write(f"Original Sentences: {metadata.get('orig_sents', 0)} | Summary Sentences: {metadata.get('sum_sents', 0)}\n")
            f.write(f"Generated At: {metadata.get('timestamp', 'N/A')}\n\n")
            
            f.write("-" * 60 + "\n")
            f.write("EXTRACTED SUMMARY:\n")
            f.write("-" * 60 + "\n\n")
            f.write(summary_text)
            f.write("\n\n")
            
            keywords = metadata.get('keywords', [])
            if keywords:
                f.write("-" * 60 + "\n")
                f.write("KEY TERMS:\n")
                f.write("-" * 60 + "\n")
                f.write(", ".join(keywords) + "\n")
                
        return output_filepath
    except Exception as e:
        raise RuntimeError(f"Failed to generate TXT document: {e}")
